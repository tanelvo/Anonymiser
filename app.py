from docx import Document
from flask import Flask, request, jsonify
from flask_cors import CORS
from name_entity_recognizer import process_document
from anonymiser import anonymize_text

app = Flask(__name__)
CORS(app)

def extract_text_with_breaks(doc):
    text = ""
    line_break_positions = []
    page_break_positions = []

    for paragraph in doc.paragraphs:
        text += paragraph.text
        # Track line breaks based on paragraphs and "\n" symbols
        if paragraph.text.strip():
            line_break_positions.append(len(text))
        text += "\n"

    # Check for page breaks in document structure
    for i, element in enumerate(doc.element.body):
        if element.tag.endswith('sectPr'):
            page_break_positions.append(len(text))

    return text, line_break_positions, page_break_positions

# Function to re-insert line and page breaks after processing
def restore_breaks(text, line_break_positions, page_break_positions):
    restored_text = list(text)
    
    for pos in line_break_positions:
        restored_text.insert(pos, '\n')

    for pos in page_break_positions:
        restored_text.insert(pos, '\f')  #'\f' for page breaks
    
    return ''.join(restored_text)


@app.route('/file/', methods = ['POST'])
def file_process():
    file = request.files['file']
    doc = Document(file)
    text = "\n".join([paragraph.text.replace('\r', '\n') for paragraph in doc.paragraphs])
    response_data = process_document(text)
    return jsonify(response_data)


@app.route('/text/', methods=['POST'])
def text_process():
    data = request.json
    print("========================================")
    print(data)
    text = data.get('text', '')
    response_data = process_document(text)
    return jsonify(response_data)

@app.route('/anonymise/', methods=['POST'])
def anonymize():
    text = anonymize_text(request.json)
    return jsonify(text)
    
if __name__ == '__main__':
    app.run(host='localhost', port=8080)
